import io
import os
import re
import pickle
import mimetypes
import zipfile
from datetime import datetime

import pandas as pd
import streamlit as st

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaIoBaseUpload


# ============================================================
# CONFIGURACIÓN
# ============================================================

st.set_page_config(
    page_title="Control de CV y Probatorios - Dra. Günther",
    page_icon="📄",
    layout="wide",
)

CATEGORIAS = [
    "Coordinación de Libros",
    "Capítulos de Libros / Artículos",
    "Ponencias y Conferencias",
    "Dirección de Tesis y Sinodalías",
    "Proyectos de Investigación",
    "Organización de Eventos Académicos",
    "Arbitraje y Trabajo Editorial (Dictaminación)",
    "Gestión Académico-Administrativa",
    "Difusión y Divulgación",
    "Diseño de Planes y Programas de Estudio",
    "Presentaciones de Libros",
    "Comisiones",
    "Cursos e Impartición de Clases",
    "Premios y Reconocimientos",
    "Asesorías",
]

NOMBRE_CARPETA_RAIZ = "CV — Sistema de Gestión"
NOMBRE_EXCEL = "Base_de_Datos_Probatorios_y_CV.xlsx"

ESTRUCTURA_CARPETAS = [
    "00 — Administración",
    "01 — Datos personales y CV",
    "02 — Probatorios",
    "03 — Formación académica",
    "04 — Docencia",
    "05 — Investigación",
    "06 — Ponencias y congresos",
    "07 — Publicaciones",
    "08 — Gestión y cargos",
    "09 — Reconocimientos",
    "10 — CV generados",
]

ANIOS_PROBATORIOS = list(range(2020, datetime.now().year + 2))

HOJA_ACTIVIDADES = "Actividades"
HOJA_PROBATORIOS = "Probatorios"
HOJA_LEGACY = "Base de datos"

COLUMNAS_ACTIVIDADES = [
    "ID",
    "Año",
    "Fecha",
    "Categoría",
    "Rol",
    "Título",
    "Institución",
    "Lugar",
    "Estado_Probatorio",
    "Incluir_en_CV",
    "Detalle_CV",
    "Notas_Observaciones",
]

COLUMNAS_PROBATORIOS = [
    "ID_Probatorio",
    "ID_Actividad",
    "Orden",
    "Nombre_Archivo",
    "Enlace_Drive_Probatorio",
    "ID_Drive_Probatorio",
    "Tipo_Archivo",
    "Fecha_Registro",
]

COLUMNAS_OBSOLETAS = [
    "Componente_SNII",
    "Tipo_Producto_SNII",
    "Categoría_CV",
    "Rol_Participación",
    "Título_Actividad_o_Publicación",
    "Evento_Revista_Libro",
    "Institución_Organización",
    "Modalidad",
    "Autores",
    "Coautores",
    "Nivel_Formación",
    "Estudiantes_Beneficiados",
    "Proyecto_Línea_Investigación",
    "Descripción_Aportación",
    "Impacto_Beneficio_Social",
    "Características_SNII",
    "Arbitrado",
    "Publicado",
    "Revista_Editorial",
    "Volumen_Número",
    "Páginas",
    "ISBN_ISSN",
    "DOI_URL",
    "Incluir_en_CV_SNII",
    "Redacción",
]

SEPARADOR_ARCHIVOS = " || "


# ============================================================
# GOOGLE DRIVE
# ============================================================

@st.cache_resource
def obtener_servicio_drive():
    """Obtiene el servicio autenticado de Google Drive."""

    creds = None

    if os.path.exists("token.pickle"):
        with open("token.pickle", "rb") as token:
            creds = pickle.load(token)

    if creds and creds.expired and creds.refresh_token:
        try:
            creds.refresh(Request())
        except Exception as e:
            st.error(f"Error al refrescar las credenciales de Google: {e}")
            return None

    if not creds:
        st.error("⚠️ No se encontró el archivo 'token.pickle'.")
        return None

    try:
        return build("drive", "v3", credentials=creds)
    except Exception as e:
        st.error(f"Error al crear el servicio de Google Drive: {e}")
        return None


# ============================================================
# CARPETAS
# ============================================================

def escapar_query_drive(texto):
    return str(texto).replace("\\", "\\\\").replace("'", "\\'")


def buscar_carpeta(service, nombre, parent_id=None):
    try:
        nombre_escapado = escapar_query_drive(nombre)

        if parent_id:
            query = (
                f"name = '{nombre_escapado}' "
                "and mimeType = 'application/vnd.google-apps.folder' "
                "and trashed = false "
                f"and '{parent_id}' in parents"
            )
        else:
            query = (
                f"name = '{nombre_escapado}' "
                "and mimeType = 'application/vnd.google-apps.folder' "
                "and trashed = false "
                "and 'root' in parents"
            )

        resultado = service.files().list(
            q=query,
            spaces="drive",
            fields="files(id,name,webViewLink,parents)",
            pageSize=10,
        ).execute()

        carpetas = resultado.get("files", [])
        return carpetas[0] if carpetas else None

    except Exception as e:
        st.error(f"Error al buscar la carpeta '{nombre}': {e}")
        return None


def crear_carpeta(service, nombre, parent_id=None):
    try:
        metadata = {
            "name": nombre,
            "mimeType": "application/vnd.google-apps.folder",
        }

        if parent_id:
            metadata["parents"] = [parent_id]

        return service.files().create(
            body=metadata,
            fields="id,name,webViewLink,parents",
        ).execute()

    except Exception as e:
        st.error(f"Error al crear la carpeta '{nombre}': {e}")
        return None


def obtener_o_crear_carpeta(service, nombre, parent_id=None):
    carpeta = buscar_carpeta(service, nombre, parent_id)
    if carpeta:
        return carpeta
    return crear_carpeta(service, nombre, parent_id)


@st.cache_resource
def inicializar_estructura_drive(_service):
    estructura = {}

    raiz = obtener_o_crear_carpeta(_service, NOMBRE_CARPETA_RAIZ)
    if not raiz:
        return {}

    estructura["raiz"] = raiz

    for nombre in ESTRUCTURA_CARPETAS:
        carpeta = obtener_o_crear_carpeta(_service, nombre, raiz["id"])
        if carpeta:
            estructura[nombre] = carpeta

    carpeta_probatorios = estructura.get("02 — Probatorios")
    if not carpeta_probatorios:
        return estructura

    estructura["probatorios"] = carpeta_probatorios

    for anio in ANIOS_PROBATORIOS:
        carpeta_anio = obtener_o_crear_carpeta(
            _service,
            str(anio),
            carpeta_probatorios["id"],
        )

        if not carpeta_anio:
            continue

        estructura[f"probatorios_{anio}"] = carpeta_anio

        for categoria in CATEGORIAS:
            carpeta_categoria = obtener_o_crear_carpeta(
                _service,
                categoria,
                carpeta_anio["id"],
            )

            if carpeta_categoria:
                estructura[f"probatorios_{anio}_{categoria}"] = carpeta_categoria

    return estructura


def obtener_carpeta_probatorio(service, estructura, anio, categoria):
    clave = f"probatorios_{anio}_{categoria}"

    if clave in estructura:
        return estructura[clave]

    carpeta_probatorios = estructura.get("probatorios")
    if not carpeta_probatorios:
        return None

    carpeta_anio = obtener_o_crear_carpeta(
        service,
        str(anio),
        carpeta_probatorios["id"],
    )
    if not carpeta_anio:
        return None

    estructura[f"probatorios_{anio}"] = carpeta_anio

    carpeta_categoria = obtener_o_crear_carpeta(
        service,
        categoria,
        carpeta_anio["id"],
    )

    if carpeta_categoria:
        estructura[clave] = carpeta_categoria

    return carpeta_categoria


# ============================================================
# ARCHIVOS DRIVE
# ============================================================

def obtener_mimetype(nombre_archivo):
    mimetype, _ = mimetypes.guess_type(nombre_archivo)
    return mimetype or "application/octet-stream"


def construir_url_drive(archivo_id):
    archivo_id = str(archivo_id or "").strip()
    if not archivo_id:
        return ""
    return f"https://drive.google.com/file/d/{archivo_id}/view"


def extraer_drive_id_de_url(url):
    texto = str(url or "").strip()
    if not texto:
        return ""

    patrones = [
        r"/file/d/([A-Za-z0-9_-]+)",
        r"/d/([A-Za-z0-9_-]+)",
        r"[?&]id=([A-Za-z0-9_-]+)",
    ]

    for patron in patrones:
        match = re.search(patron, texto)
        if match:
            return match.group(1)

    return ""


def obtener_metadata_drive(service, archivo_id):
    if not archivo_id:
        return None

    try:
        return service.files().get(
            fileId=str(archivo_id).strip(),
            fields="id,name,mimeType,webViewLink,parents,trashed",
        ).execute()
    except Exception:
        return None


def obtener_enlace_drive(service, archivo_id, enlace_previo=""):
    archivo_id = str(archivo_id or "").strip()
    enlace_previo = str(enlace_previo or "").strip()

    if not archivo_id:
        return enlace_previo

    metadata = obtener_metadata_drive(service, archivo_id)
    if metadata and metadata.get("webViewLink"):
        return metadata["webViewLink"]

    return enlace_previo or construir_url_drive(archivo_id)


def subir_a_google_drive(service, nombre_archivo, bytes_archivo, carpeta_destino):
    if not carpeta_destino:
        st.error("No se encontró la carpeta de destino en Drive.")
        return None

    try:
        metadata = {
            "name": nombre_archivo,
            "parents": [carpeta_destino["id"]],
        }

        media = MediaIoBaseUpload(
            io.BytesIO(bytes_archivo),
            mimetype=obtener_mimetype(nombre_archivo),
            resumable=True,
        )

        archivo = service.files().create(
            body=metadata,
            media_body=media,
            fields="id,name,webViewLink,parents,mimeType",
        ).execute()

        archivo_id = archivo.get("id", "")
        enlace = archivo.get("webViewLink", "")

        # Garantía adicional: si create no devolvió webViewLink,
        # lo recuperamos con files.get y, como último recurso,
        # construimos la URL estándar de Drive.
        if archivo_id and not enlace:
            enlace = obtener_enlace_drive(service, archivo_id)

        return {
            "nombre": archivo.get("name", nombre_archivo),
            "enlace": enlace or construir_url_drive(archivo_id),
            "id": archivo_id,
            "mimetype": archivo.get("mimeType") or obtener_mimetype(nombre_archivo),
        }

    except Exception as e:
        st.error(f"Error al subir '{nombre_archivo}': {e}")
        return None


def obtener_archivo_drive(service, archivo_id):
    try:
        metadata = service.files().get(
            fileId=archivo_id,
            fields="id,name,mimeType",
        ).execute()

        request = service.files().get_media(fileId=archivo_id)
        buffer = io.BytesIO()
        downloader = MediaIoBaseDownload(buffer, request)

        done = False
        while not done:
            _, done = downloader.next_chunk()

        buffer.seek(0)
        return metadata, buffer.getvalue()

    except Exception as e:
        st.error(f"No se pudo descargar el archivo {archivo_id}: {e}")
        return None, None


def eliminar_archivo_drive(service, archivo_id):
    if not archivo_id:
        return True

    try:
        service.files().update(
            fileId=str(archivo_id).strip(),
            body={"trashed": True},
        ).execute()
        return True

    except Exception as e:
        st.error(f"Error al enviar archivo a papelera: {e}")
        return False


def mover_archivo_drive(service, archivo_id, carpeta_destino):
    if not archivo_id or not carpeta_destino:
        return False

    try:
        archivo = service.files().get(
            fileId=str(archivo_id).strip(),
            fields="parents",
        ).execute()

        padres_actuales = archivo.get("parents", [])

        kwargs = {
            "fileId": str(archivo_id).strip(),
            "addParents": carpeta_destino["id"],
            "fields": "id,parents",
        }

        if padres_actuales:
            kwargs["removeParents"] = ",".join(padres_actuales)

        service.files().update(**kwargs).execute()
        return True

    except Exception as e:
        st.error(f"Error al mover archivo: {e}")
        return False


def renombrar_archivo_drive(service, archivo_id, nuevo_nombre):
    if not archivo_id:
        return False

    try:
        service.files().update(
            fileId=str(archivo_id).strip(),
            body={"name": nuevo_nombre},
        ).execute()
        return True

    except Exception as e:
        st.error(f"Error al renombrar archivo: {e}")
        return False


# ============================================================
# NOMBRES
# ============================================================

def limpiar_nombre_archivo(texto):
    texto = str(texto or "")
    texto = re.sub(r'[\\/:*?"<>|]+', "_", texto)
    texto = re.sub(r"\s+", " ", texto).strip()
    return texto


def nombre_probatorio(titulo, anio, categoria, numero, extension):
    titulo_limpio = limpiar_nombre_archivo(titulo)[:100]
    categoria_limpia = limpiar_nombre_archivo(categoria).replace(" ", "_")
    extension = str(extension or "").lower()

    return (
        f"{anio}_"
        f"{categoria_limpia}_"
        f"{titulo_limpio}_"
        f"Probatorio_{numero}"
        f"{extension}"
    )


# ============================================================
# EXCEL: LOCALIZACIÓN, LECTURA, ESCRITURA Y RESPALDO
# ============================================================

def buscar_excel_en_drive(service):
    """Busca por nombre EXACTO para no confundir el archivo vivo con respaldos."""

    try:
        nombre = escapar_query_drive(NOMBRE_EXCEL)
        resultado = service.files().list(
            q=f"name = '{nombre}' and trashed = false",
            fields="files(id,name,parents,modifiedTime)",
            orderBy="modifiedTime desc",
            pageSize=20,
        ).execute()

        archivos = resultado.get("files", [])
        if not archivos:
            return None, None

        archivo = archivos[0]
        return archivo["id"], archivo["name"]

    except Exception as e:
        st.error(f"Error al buscar Excel: {e}")
        return None, None


def descargar_excel_drive(service, file_id):
    try:
        request = service.files().get_media(fileId=file_id)
        buffer = io.BytesIO()
        downloader = MediaIoBaseDownload(buffer, request)

        done = False
        while not done:
            _, done = downloader.next_chunk()

        buffer.seek(0)
        return buffer.getvalue()

    except Exception as e:
        st.error(f"Error al descargar el Excel: {e}")
        return None


def leer_libro_excel(excel_bytes):
    if not excel_bytes:
        return {}, []

    buffer = io.BytesIO(excel_bytes)

    try:
        with pd.ExcelFile(buffer) as xls:
            nombres = xls.sheet_names
            hojas = {
                nombre: pd.read_excel(xls, sheet_name=nombre)
                for nombre in nombres
            }
        return hojas, nombres

    except Exception as e:
        st.error(f"No se pudo leer el archivo Excel: {e}")
        return {}, []


def ajustar_hoja_excel(ws, anchos=None, columna_hipervinculo=None):
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions

    if anchos:
        for columna, ancho in anchos.items():
            ws.column_dimensions[columna].width = ancho

    if columna_hipervinculo:
        encabezados = {
            cell.value: cell.column
            for cell in ws[1]
        }
        numero_columna = encabezados.get(columna_hipervinculo)

        if numero_columna:
            for fila in range(2, ws.max_row + 1):
                celda = ws.cell(row=fila, column=numero_columna)
                valor = str(celda.value or "").strip()
                if valor.startswith("http://") or valor.startswith("https://"):
                    celda.hyperlink = valor
                    celda.style = "Hyperlink"


def construir_excel_bytes(df_actividades, df_probatorios):
    output = io.BytesIO()

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df_actividades[COLUMNAS_ACTIVIDADES].to_excel(
            writer,
            index=False,
            sheet_name=HOJA_ACTIVIDADES,
        )

        df_probatorios[COLUMNAS_PROBATORIOS].to_excel(
            writer,
            index=False,
            sheet_name=HOJA_PROBATORIOS,
        )

        ws_act = writer.sheets[HOJA_ACTIVIDADES]
        ws_prob = writer.sheets[HOJA_PROBATORIOS]

        ajustar_hoja_excel(
            ws_act,
            anchos={
                "A": 14,
                "B": 10,
                "C": 14,
                "D": 38,
                "E": 28,
                "F": 55,
                "G": 40,
                "H": 28,
                "I": 24,
                "J": 16,
                "K": 70,
                "L": 55,
            },
        )

        ajustar_hoja_excel(
            ws_prob,
            anchos={
                "A": 16,
                "B": 14,
                "C": 10,
                "D": 70,
                "E": 70,
                "F": 44,
                "G": 30,
                "H": 22,
            },
            columna_hipervinculo="Enlace_Drive_Probatorio",
        )

    output.seek(0)
    return output.getvalue()


def actualizar_excel_drive(service, file_id, df_actividades, df_probatorios):
    try:
        excel_bytes = construir_excel_bytes(df_actividades, df_probatorios)

        media = MediaIoBaseUpload(
            io.BytesIO(excel_bytes),
            mimetype=(
                "application/vnd.openxmlformats-officedocument."
                "spreadsheetml.sheet"
            ),
            resumable=True,
        )

        service.files().update(
            fileId=file_id,
            media_body=media,
        ).execute()

        return True

    except Exception as e:
        st.error(f"No se pudo actualizar el Excel en Drive: {e}")
        return False


def crear_respaldo_excel(service, file_id):
    """Crea una copia del Excel antes de la migración estructural."""

    try:
        metadata = service.files().get(
            fileId=file_id,
            fields="name,parents",
        ).execute()

        nombre_actual = metadata.get("name", NOMBRE_EXCEL)
        base, extension = os.path.splitext(nombre_actual)
        marca = datetime.now().strftime("%Y%m%d_%H%M%S")
        nombre_respaldo = f"{base}_RESPALDO_{marca}{extension or '.xlsx'}"

        cuerpo = {"name": nombre_respaldo}
        padres = metadata.get("parents", [])
        if padres:
            cuerpo["parents"] = padres

        copia = service.files().copy(
            fileId=file_id,
            body=cuerpo,
            fields="id,name",
        ).execute()

        return copia

    except Exception as e:
        st.error(f"No se pudo crear el respaldo del Excel: {e}")
        return None


# ============================================================
# NORMALIZACIÓN DE DATAFRAMES
# ============================================================

def limpiar_valor(valor, default=""):
    if valor is None:
        return default
    try:
        if pd.isna(valor):
            return default
    except Exception:
        pass
    return valor


def preparar_actividades(df):
    df = df.copy() if df is not None else pd.DataFrame()

    equivalencias = {
        "Categoría_CV": "Categoría",
        "Rol_Participación": "Rol",
        "Título_Actividad_o_Publicación": "Título",
        "Institución_Organización": "Institución",
    }

    for antigua, nueva in equivalencias.items():
        if antigua in df.columns and nueva not in df.columns:
            df[nueva] = df[antigua]

    for columna in COLUMNAS_ACTIVIDADES:
        if columna not in df.columns:
            df[columna] = ""

    df = df[COLUMNAS_ACTIVIDADES].copy()

    for columna in df.columns:
        df[columna] = df[columna].apply(lambda x: limpiar_valor(x, ""))

    if "ID" in df.columns:
        df["ID"] = df["ID"].astype(str).str.strip()

    return df


def preparar_probatorios(df, service=None):
    df = df.copy() if df is not None else pd.DataFrame()

    # Compatibilidad con nombres alternativos por si la hoja ya existía.
    equivalencias = {
        "Archivo": "Nombre_Archivo",
        "Drive_ID": "ID_Drive_Probatorio",
        "Enlace": "Enlace_Drive_Probatorio",
    }

    for antigua, nueva in equivalencias.items():
        if antigua in df.columns and nueva not in df.columns:
            df[nueva] = df[antigua]

    for columna in COLUMNAS_PROBATORIOS:
        if columna not in df.columns:
            df[columna] = ""

    df = df[COLUMNAS_PROBATORIOS].copy()

    for columna in df.columns:
        df[columna] = df[columna].apply(lambda x: limpiar_valor(x, ""))

    # Completar ID de Drive desde URL cuando sea posible.
    for indice, fila in df.iterrows():
        drive_id = str(fila.get("ID_Drive_Probatorio", "") or "").strip()
        enlace = str(fila.get("Enlace_Drive_Probatorio", "") or "").strip()

        if not drive_id and enlace:
            drive_id = extraer_drive_id_de_url(enlace)
            if drive_id:
                df.at[indice, "ID_Drive_Probatorio"] = drive_id

        if drive_id and not enlace:
            if service:
                enlace = obtener_enlace_drive(service, drive_id)
            else:
                enlace = construir_url_drive(drive_id)
            df.at[indice, "Enlace_Drive_Probatorio"] = enlace

    # Orden: si está vacío, lo reconstruimos por actividad.
    for actividad_id, indices in df.groupby("ID_Actividad", sort=False).groups.items():
        if not str(actividad_id or "").strip():
            continue

        contador = 1
        for indice in list(indices):
            valor_orden = df.at[indice, "Orden"]
            try:
                orden = int(float(valor_orden))
            except Exception:
                orden = contador
                df.at[indice, "Orden"] = orden
            contador = max(contador + 1, orden + 1)

    return df


def convertir_celda_a_lista(valor):
    if valor is None:
        return []

    try:
        if pd.isna(valor):
            return []
    except Exception:
        pass

    texto = str(valor).strip()
    if not texto:
        return []

    valores = [x.strip() for x in texto.split(SEPARADOR_ARCHIVOS)]

    return [
        x
        for x in valores
        if x
        and x.lower() not in {"nan", "none", "sin_pdf", "sin_enlace"}
    ]


def siguiente_id_registro(df):
    numeros = []

    if "ID" in df.columns:
        for valor in df["ID"].dropna():
            match = re.fullmatch(r"ACT-(\d+)", str(valor).strip(), re.IGNORECASE)
            if match:
                numeros.append(int(match.group(1)))

    siguiente = max(numeros, default=0) + 1
    return f"ACT-{siguiente:03d}"


def siguiente_numero_probatorio(df_probatorios):
    numeros = []

    if "ID_Probatorio" in df_probatorios.columns:
        for valor in df_probatorios["ID_Probatorio"].dropna():
            match = re.fullmatch(r"PROB-(\d+)", str(valor).strip(), re.IGNORECASE)
            if match:
                numeros.append(int(match.group(1)))

    return max(numeros, default=0) + 1


def generar_ids_probatorios(df_probatorios, cantidad):
    inicio = siguiente_numero_probatorio(df_probatorios)
    return [f"PROB-{n:04d}" for n in range(inicio, inicio + cantidad)]


def migrar_legacy_a_normalizado(df_legacy, service=None):
    """Convierte una fila por actividad + celdas concatenadas a dos tablas 1:N."""

    legacy = df_legacy.copy()

    equivalencias = {
        "Categoría_CV": "Categoría",
        "Rol_Participación": "Rol",
        "Título_Actividad_o_Publicación": "Título",
        "Institución_Organización": "Institución",
    }

    for antigua, nueva in equivalencias.items():
        if antigua in legacy.columns and nueva not in legacy.columns:
            legacy[nueva] = legacy[antigua]

    for columna in COLUMNAS_ACTIVIDADES:
        if columna not in legacy.columns:
            legacy[columna] = ""

    df_actividades = preparar_actividades(legacy)

    filas_probatorios = []
    contador_probatorio = 1

    for _, fila in legacy.iterrows():
        actividad_id = str(limpiar_valor(fila.get("ID", ""), "")).strip()
        if not actividad_id:
            continue

        nombres = convertir_celda_a_lista(fila.get("Nombre_Archivo_PDF", ""))
        enlaces = convertir_celda_a_lista(fila.get("Enlace_Drive_Probatorio", ""))
        ids = convertir_celda_a_lista(fila.get("ID_Drive_Probatorio", ""))

        cantidad = max(len(nombres), len(enlaces), len(ids))

        for i in range(cantidad):
            nombre = nombres[i] if i < len(nombres) else ""
            enlace = enlaces[i] if i < len(enlaces) else ""
            drive_id = ids[i] if i < len(ids) else ""

            if not drive_id and enlace:
                drive_id = extraer_drive_id_de_url(enlace)

            if drive_id and not enlace:
                if service:
                    enlace = obtener_enlace_drive(service, drive_id)
                else:
                    enlace = construir_url_drive(drive_id)

            mimetype = obtener_mimetype(nombre) if nombre else ""

            filas_probatorios.append(
                {
                    "ID_Probatorio": f"PROB-{contador_probatorio:04d}",
                    "ID_Actividad": actividad_id,
                    "Orden": i + 1,
                    "Nombre_Archivo": nombre,
                    "Enlace_Drive_Probatorio": enlace,
                    "ID_Drive_Probatorio": drive_id,
                    "Tipo_Archivo": mimetype,
                    "Fecha_Registro": "",
                }
            )

            contador_probatorio += 1

    df_probatorios = pd.DataFrame(filas_probatorios, columns=COLUMNAS_PROBATORIOS)
    df_probatorios = preparar_probatorios(df_probatorios, service=service)

    return df_actividades, df_probatorios


def reparar_urls_probatorios(service, df_probatorios):
    actualizado = df_probatorios.copy()
    cambios = 0

    for indice, fila in actualizado.iterrows():
        drive_id = str(fila.get("ID_Drive_Probatorio", "") or "").strip()
        enlace = str(fila.get("Enlace_Drive_Probatorio", "") or "").strip()

        if not drive_id and enlace:
            drive_id_extraido = extraer_drive_id_de_url(enlace)
            if drive_id_extraido:
                actualizado.at[indice, "ID_Drive_Probatorio"] = drive_id_extraido
                drive_id = drive_id_extraido
                cambios += 1

        if drive_id and not enlace:
            nuevo_enlace = obtener_enlace_drive(service, drive_id)
            if nuevo_enlace:
                actualizado.at[indice, "Enlace_Drive_Probatorio"] = nuevo_enlace
                cambios += 1

    return actualizado, cambios


def validar_integridad(df_actividades, df_probatorios):
    ids_actividades = set(df_actividades["ID"].astype(str).str.strip())
    ids_probatorios = df_probatorios["ID_Probatorio"].astype(str).str.strip()

    duplicados_actividades = int(
        df_actividades["ID"].astype(str).str.strip().duplicated().sum()
    )
    duplicados_probatorios = int(ids_probatorios.duplicated().sum())

    huerfanos = df_probatorios[
        ~df_probatorios["ID_Actividad"].astype(str).str.strip().isin(ids_actividades)
    ]

    sin_drive_id = df_probatorios[
        df_probatorios["ID_Drive_Probatorio"].astype(str).str.strip() == ""
    ]

    sin_url = df_probatorios[
        df_probatorios["Enlace_Drive_Probatorio"].astype(str).str.strip() == ""
    ]

    return {
        "duplicados_actividades": duplicados_actividades,
        "duplicados_probatorios": duplicados_probatorios,
        "huerfanos": len(huerfanos),
        "sin_drive_id": len(sin_drive_id),
        "sin_url": len(sin_url),
    }


# ============================================================
# UTILIDADES DE ACTIVIDADES / PROBATORIOS
# ============================================================

def valor_fila(fila, columna, default=""):
    try:
        valor = fila[columna]
    except Exception:
        return default

    try:
        if pd.isna(valor):
            return default
    except Exception:
        pass

    return valor


def normalizar_anio(valor, default=None):
    if default is None:
        default = datetime.now().year

    try:
        return int(float(valor))
    except Exception:
        return int(default)


def normalizar_fecha(valor):
    try:
        fecha = pd.to_datetime(valor, errors="coerce")
        if pd.isna(fecha):
            return datetime.now().date()
        return fecha.date()
    except Exception:
        return datetime.now().date()


def formatear_fecha(valor):
    try:
        fecha = pd.to_datetime(valor, errors="coerce")
        if pd.isna(fecha):
            return ""
        return fecha.strftime("%d/%m/%Y")
    except Exception:
        return str(valor or "")


def probatorios_de_actividad(df_probatorios, actividad_id):
    resultado = df_probatorios[
        df_probatorios["ID_Actividad"].astype(str).str.strip()
        == str(actividad_id).strip()
    ].copy()

    if resultado.empty:
        return resultado

    resultado["_orden_num"] = pd.to_numeric(resultado["Orden"], errors="coerce")
    resultado = resultado.sort_values(
        ["_orden_num", "ID_Probatorio"],
        na_position="last",
    ).drop(columns=["_orden_num"])

    return resultado


def reordenar_probatorios_actividad(df_probatorios, actividad_id):
    actualizado = df_probatorios.copy()
    indices = probatorios_de_actividad(actualizado, actividad_id).index.tolist()

    for orden, indice in enumerate(indices, start=1):
        actualizado.at[indice, "Orden"] = orden

    return actualizado


def sincronizar_nombres_y_carpeta(
    service,
    estructura_drive,
    df_probatorios,
    actividad_id,
    titulo,
    anio,
    categoria,
):
    actualizado = reordenar_probatorios_actividad(df_probatorios, actividad_id)
    sub = probatorios_de_actividad(actualizado, actividad_id)

    carpeta_destino = obtener_carpeta_probatorio(
        service,
        estructura_drive,
        anio,
        categoria,
    )

    for orden, (indice, fila) in enumerate(sub.iterrows(), start=1):
        drive_id = str(fila.get("ID_Drive_Probatorio", "") or "").strip()
        nombre_actual = str(fila.get("Nombre_Archivo", "") or "").strip()

        if not drive_id:
            continue

        extension = os.path.splitext(nombre_actual)[1].lower()
        if not extension:
            metadata = obtener_metadata_drive(service, drive_id)
            if metadata:
                extension = os.path.splitext(metadata.get("name", ""))[1].lower()

        nuevo_nombre = nombre_probatorio(
            titulo,
            anio,
            categoria,
            orden,
            extension,
        )

        mover_archivo_drive(service, drive_id, carpeta_destino)

        if nuevo_nombre and nuevo_nombre != nombre_actual:
            if renombrar_archivo_drive(service, drive_id, nuevo_nombre):
                actualizado.at[indice, "Nombre_Archivo"] = nuevo_nombre

        enlace = obtener_enlace_drive(
            service,
            drive_id,
            fila.get("Enlace_Drive_Probatorio", ""),
        )
        actualizado.at[indice, "Enlace_Drive_Probatorio"] = enlace
        actualizado.at[indice, "Orden"] = orden

    return actualizado


def crear_filas_probatorios_subidos(
    service,
    estructura_drive,
    archivos,
    df_probatorios,
    actividad_id,
    titulo,
    anio,
    categoria,
    orden_inicial,
):
    archivos = archivos or []
    if not archivos:
        return [], []

    nuevos_ids = generar_ids_probatorios(df_probatorios, len(archivos))
    filas = []
    ids_drive_subidos = []

    carpeta = obtener_carpeta_probatorio(
        service,
        estructura_drive,
        anio,
        categoria,
    )

    for posicion, archivo in enumerate(archivos):
        orden = orden_inicial + posicion
        extension = os.path.splitext(archivo.name)[1].lower()

        nombre = nombre_probatorio(
            titulo,
            anio,
            categoria,
            orden,
            extension,
        )

        resultado = subir_a_google_drive(
            service,
            nombre,
            archivo.getvalue(),
            carpeta,
        )

        if not resultado or not resultado.get("id"):
            for drive_id in ids_drive_subidos:
                eliminar_archivo_drive(service, drive_id)
            return None, []

        ids_drive_subidos.append(resultado["id"])

        filas.append(
            {
                "ID_Probatorio": nuevos_ids[posicion],
                "ID_Actividad": actividad_id,
                "Orden": orden,
                "Nombre_Archivo": resultado["nombre"],
                "Enlace_Drive_Probatorio": resultado["enlace"],
                "ID_Drive_Probatorio": resultado["id"],
                "Tipo_Archivo": resultado["mimetype"],
                "Fecha_Registro": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            }
        )

    return filas, ids_drive_subidos


# ============================================================
# ZIP
# ============================================================

def crear_zip_probatorios(service, df_actividades, df_probatorios):
    buffer_zip = io.BytesIO()
    archivos_agregados = 0
    nombres_zip = set()

    ids_seleccionados = set(df_actividades["ID"].astype(str).str.strip())
    probatorios = df_probatorios[
        df_probatorios["ID_Actividad"].astype(str).str.strip().isin(ids_seleccionados)
    ].copy()

    with zipfile.ZipFile(buffer_zip, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for _, probatorio in probatorios.iterrows():
            actividad_id = str(probatorio.get("ID_Actividad", "SIN_ID")).strip()
            archivo_id = str(probatorio.get("ID_Drive_Probatorio", "") or "").strip()
            orden = normalizar_anio(probatorio.get("Orden", 1), default=1)

            if not archivo_id:
                continue

            metadata, contenido = obtener_archivo_drive(service, archivo_id)
            if not contenido:
                continue

            nombre_original = (
                metadata.get("name", "")
                if metadata
                else str(probatorio.get("Nombre_Archivo", "") or "")
            )
            nombre_original = limpiar_nombre_archivo(nombre_original or f"probatorio_{orden}")

            nombre_zip = f"{actividad_id}/{orden:02d}_{nombre_original}"
            nombre_base = nombre_zip
            contador = 2

            while nombre_zip in nombres_zip:
                raiz, ext = os.path.splitext(nombre_base)
                nombre_zip = f"{raiz}_{contador}{ext}"
                contador += 1

            nombres_zip.add(nombre_zip)
            zip_file.writestr(nombre_zip, contenido)
            archivos_agregados += 1

    buffer_zip.seek(0)
    return buffer_zip.getvalue(), archivos_agregados


# ============================================================
# WORD
# ============================================================

def crear_cv_word(df_actividades):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DRA. MARÍA GRISELDA GÜNTHER")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 51, 102)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("CURRÍCULUM VITAE — SÍNTESIS EJECUTIVA")
    run2.font.size = Pt(10.5)
    run2.font.italic = True

    df_cv = df_actividades[
        df_actividades["Incluir_en_CV"]
        .astype(str)
        .str.strip()
        .str.lower()
        == "sí"
    ].copy()

    if df_cv.empty:
        doc.add_paragraph("No existen actividades marcadas para incluir en el CV.")
    else:
        df_cv["Año_num"] = pd.to_numeric(df_cv["Año"], errors="coerce")
        df_cv = df_cv.sort_values(["Año_num", "Fecha"], ascending=[False, False])

        categorias_presentes = set(df_cv["Categoría"].dropna().astype(str))

        for categoria in CATEGORIAS:
            if categoria not in categorias_presentes:
                continue

            sub_df = df_cv[df_cv["Categoría"] == categoria]
            if sub_df.empty:
                continue

            p_cat = doc.add_paragraph()
            p_cat.paragraph_format.space_before = Pt(14)
            p_cat.paragraph_format.space_after = Pt(6)

            run_cat = p_cat.add_run(categoria)
            run_cat.bold = True
            run_cat.font.size = Pt(12.5)
            run_cat.font.color.rgb = RGBColor(0, 51, 102)

            for _, fila in sub_df.iterrows():
                titulo = str(valor_fila(fila, "Título", "")).strip()
                rol = str(valor_fila(fila, "Rol", "")).strip()
                institucion = str(valor_fila(fila, "Institución", "")).strip()
                lugar = str(valor_fila(fila, "Lugar", "")).strip()
                fecha = formatear_fecha(valor_fila(fila, "Fecha", ""))
                detalle_cv = str(valor_fila(fila, "Detalle_CV", "")).strip()

                if not titulo:
                    continue

                p_item = doc.add_paragraph(style="List Bullet")
                run_t = p_item.add_run(titulo)
                run_t.bold = True

                detalles = []
                if rol:
                    detalles.append(f"Rol: {rol}")
                if institucion:
                    detalles.append(institucion)
                if lugar:
                    detalles.append(lugar)
                if fecha:
                    detalles.append(fecha)

                if detalles:
                    p_item.add_run(". " + ", ".join(detalles) + ".")

                # Información complementaria destinada expresamente al CV.
                # Las notas internas NO se incorporan automáticamente al documento.
                if detalle_cv:
                    p_detalle = doc.add_paragraph()
                    p_detalle.paragraph_format.left_indent = Inches(0.25)
                    p_detalle.paragraph_format.space_before = Pt(0)
                    p_detalle.paragraph_format.space_after = Pt(5)
                    run_detalle = p_detalle.add_run(detalle_cv)
                    run_detalle.italic = True
                    run_detalle.font.size = Pt(10.5)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ============================================================
# APLICACIÓN
# ============================================================

st.title("📄 Sistema de Gestión de CV — Dra. María Griselda Günther")

service = obtener_servicio_drive()

if service:
    with st.spinner("🔧 Verificando estructura de Google Drive..."):
        estructura_drive = inicializar_estructura_drive(service)

    if not estructura_drive:
        st.error("No fue posible inicializar Google Drive.")
        st.stop()

    # --------------------------------------------------------
    # LOCALIZAR / SUBIR EXCEL
    # --------------------------------------------------------

    excel_id, found_name = buscar_excel_en_drive(service)

    if not excel_id:
        st.warning("No se encontró la base de datos.")

        archivo_excel = st.file_uploader(
            f"Sube {NOMBRE_EXCEL}",
            type=["xlsx"],
        )

        if archivo_excel:
            metadata = {"name": NOMBRE_EXCEL}

            carpeta_admin = estructura_drive.get("00 — Administración")
            if carpeta_admin:
                metadata["parents"] = [carpeta_admin["id"]]

            media = MediaIoBaseUpload(
                io.BytesIO(archivo_excel.getvalue()),
                mimetype=(
                    "application/vnd.openxmlformats-officedocument."
                    "spreadsheetml.sheet"
                ),
                resumable=True,
            )

            service.files().create(
                body=metadata,
                media_body=media,
                fields="id",
            ).execute()

            st.success("Base de datos subida correctamente.")
            st.rerun()

        st.stop()

    # --------------------------------------------------------
    # LEER LIBRO Y DETECTAR ARQUITECTURA
    # --------------------------------------------------------

    excel_bytes = descargar_excel_drive(service, excel_id)
    hojas, nombres_hojas = leer_libro_excel(excel_bytes)

    if not hojas:
        st.stop()

    arquitectura_normalizada = (
        HOJA_ACTIVIDADES in hojas and HOJA_PROBATORIOS in hojas
    )

    if arquitectura_normalizada:
        columnas_actividades_existentes = set(hojas[HOJA_ACTIVIDADES].columns)

        df_actividades = preparar_actividades(hojas[HOJA_ACTIVIDADES])
        df_probatorios = preparar_probatorios(
            hojas[HOJA_PROBATORIOS],
            service=service,
        )

        # Migración menor y controlada para incorporar el campo Detalle_CV
        # a una base que ya usa la arquitectura Actividades + Probatorios.
        if "Detalle_CV" not in columnas_actividades_existentes:
            st.warning(
                "⚠️ La base ya usa Actividades + Probatorios, pero todavía no tiene "
                "el nuevo campo **Detalle_CV**. Este campo separa la información que "
                "sí debe aparecer en el CV de las Notas / Observaciones de uso interno."
            )

            if st.button(
                "🛡️ Crear respaldo y agregar campo Detalle_CV",
                type="primary",
            ):
                with st.spinner("Creando respaldo y actualizando la estructura..."):
                    respaldo = crear_respaldo_excel(service, excel_id)

                    if not respaldo:
                        st.error(
                            "La actualización se canceló porque no se pudo crear el respaldo."
                        )
                        st.stop()

                    ok = actualizar_excel_drive(
                        service,
                        excel_id,
                        df_actividades,
                        df_probatorios,
                    )

                if ok:
                    st.success(
                        f"✅ Campo Detalle_CV agregado. Respaldo creado: {respaldo['name']}"
                    )
                    st.rerun()

            st.info(
                "Esta actualización no copia automáticamente las Notas / Observaciones "
                "al nuevo campo, para evitar que información interna termine en el CV."
            )
            st.stop()

    else:
        if HOJA_LEGACY in hojas:
            hoja_origen = HOJA_LEGACY
        else:
            hoja_origen = nombres_hojas[0]

        df_legacy = hojas[hoja_origen]
        df_actividades, df_probatorios = migrar_legacy_a_normalizado(
            df_legacy,
            service=service,
        )

        st.warning(
            "⚠️ La base todavía usa la estructura histórica de una sola hoja. "
            "La aplicación ya preparó la migración a **Actividades + Probatorios**, "
            "pero no modificará el archivo hasta que pulses el botón."
        )

        col_m1, col_m2, col_m3 = st.columns(3)
        col_m1.metric("Actividades", len(df_actividades))
        col_m2.metric("Probatorios detectados", len(df_probatorios))
        col_m3.metric(
            "Probatorios con URL",
            int(
                (
                    df_probatorios["Enlace_Drive_Probatorio"]
                    .astype(str)
                    .str.strip()
                    != ""
                ).sum()
            ),
        )

        st.dataframe(
            df_probatorios.head(20),
            use_container_width=True,
            hide_index=True,
        )

        if st.button(
            "🛡️ Crear respaldo y migrar a Actividades + Probatorios",
            type="primary",
        ):
            with st.spinner("Creando respaldo y migrando la base..."):
                respaldo = crear_respaldo_excel(service, excel_id)

                if not respaldo:
                    st.error("La migración se canceló porque no se pudo crear el respaldo.")
                    st.stop()

                ok = actualizar_excel_drive(
                    service,
                    excel_id,
                    df_actividades,
                    df_probatorios,
                )

            if ok:
                st.success(
                    f"✅ Migración completada. Respaldo creado: {respaldo['name']}"
                )
                st.rerun()

        st.info(
            "Hasta completar esta migración, la aplicación permanece en modo seguro "
            "y no habilita altas, ediciones ni eliminaciones."
        )
        st.stop()

    # --------------------------------------------------------
    # REPARACIÓN DE URLS FALTANTES
    # --------------------------------------------------------

    df_prob_reparado, cantidad_reparaciones = reparar_urls_probatorios(
        service,
        df_probatorios,
    )

    if cantidad_reparaciones:
        st.warning(
            f"🔗 Se detectaron {cantidad_reparaciones} dato(s) de Drive que pueden "
            "completarse automáticamente (URL o ID)."
        )

        if st.button("🔧 Guardar reparación de enlaces de Drive"):
            if actualizar_excel_drive(
                service,
                excel_id,
                df_actividades,
                df_prob_reparado,
            ):
                st.success("Enlaces de Drive reparados y guardados.")
                st.rerun()

    # Trabajamos en memoria con la versión reparada incluso antes de persistir.
    df_probatorios = df_prob_reparado

    # ========================================================
    # TABS
    # ========================================================

    (
        tab_buscar,
        tab_editar,
        tab_nuevo,
        tab_paquetes,
        tab_cv,
    ) = st.tabs(
        [
            "🔍 Buscar",
            "✏️ Editar",
            "➕ Nueva actividad",
            "📦 Paquetes ZIP",
            "📄 Generar CV",
        ]
    )

    # ========================================================
    # BUSCAR
    # ========================================================

    with tab_buscar:
        st.subheader("🔍 Buscador de actividades")

        col1, col2, col3 = st.columns(3)

        with col1:
            anios = sorted(
                pd.to_numeric(df_actividades["Año"], errors="coerce")
                .dropna()
                .astype(int)
                .unique()
                .tolist(),
                reverse=True,
            )
            filtro_anio = st.selectbox("Año", ["Todos"] + anios)

        with col2:
            filtro_categoria = st.selectbox(
                "Categoría",
                ["Todas"] + CATEGORIAS,
            )

        with col3:
            filtro_texto = st.text_input(
                "Buscar",
                placeholder="Título, institución, ID...",
            )

        resultado = df_actividades.copy()

        if filtro_anio != "Todos":
            resultado = resultado[
                pd.to_numeric(resultado["Año"], errors="coerce")
                == int(filtro_anio)
            ]

        if filtro_categoria != "Todas":
            resultado = resultado[resultado["Categoría"] == filtro_categoria]

        if filtro_texto:
            mask = resultado.apply(
                lambda fila: fila.astype(str)
                .str.contains(
                    filtro_texto,
                    case=False,
                    na=False,
                    regex=False,
                )
                .any(),
                axis=1,
            )
            resultado = resultado[mask]

        conteo_prob = (
            df_probatorios.groupby("ID_Actividad")
            .size()
            .rename("N_Probatorios")
        )

        resultado_mostrar = resultado.merge(
            conteo_prob,
            how="left",
            left_on="ID",
            right_index=True,
        )
        resultado_mostrar["N_Probatorios"] = (
            resultado_mostrar["N_Probatorios"].fillna(0).astype(int)
        )

        st.write(f"**{len(resultado_mostrar)} registros encontrados.**")

        columnas_visibles = [
            "ID",
            "Año",
            "Fecha",
            "Categoría",
            "Rol",
            "Título",
            "Institución",
            "Lugar",
            "Estado_Probatorio",
            "Incluir_en_CV",
            "N_Probatorios",
        ]

        st.dataframe(
            resultado_mostrar[columnas_visibles],
            use_container_width=True,
            hide_index=True,
        )

        st.markdown("### 🔗 Probatorios y enlaces de Drive")

        ids_resultado = set(resultado_mostrar["ID"].astype(str))
        prob_resultado = df_probatorios[
            df_probatorios["ID_Actividad"].astype(str).isin(ids_resultado)
        ].copy()

        if prob_resultado.empty:
            st.info("No hay probatorios registrados para las actividades filtradas.")
        else:
            st.dataframe(
                prob_resultado[
                    [
                        "ID_Probatorio",
                        "ID_Actividad",
                        "Orden",
                        "Nombre_Archivo",
                        "Enlace_Drive_Probatorio",
                    ]
                ],
                use_container_width=True,
                hide_index=True,
                column_config={
                    "Enlace_Drive_Probatorio": st.column_config.LinkColumn(
                        "Probatorio en Drive",
                        display_text="Abrir en Drive",
                    )
                },
            )

        with st.expander("🩺 Diagnóstico de integridad"):
            diagnostico = validar_integridad(df_actividades, df_probatorios)
            d1, d2, d3, d4, d5 = st.columns(5)
            d1.metric("IDs actividad duplicados", diagnostico["duplicados_actividades"])
            d2.metric("IDs probatorio duplicados", diagnostico["duplicados_probatorios"])
            d3.metric("Probatorios huérfanos", diagnostico["huerfanos"])
            d4.metric("Sin Drive ID", diagnostico["sin_drive_id"])
            d5.metric("Sin URL", diagnostico["sin_url"])

    # ========================================================
    # EDITAR
    # ========================================================

    with tab_editar:
        st.subheader("✏️ Editar actividad y administrar probatorios")

        if df_actividades.empty:
            st.info("No existen registros.")
        else:
            busqueda = st.text_input(
                "Buscar registro",
                placeholder="ID, título, institución...",
                key="buscar_edicion",
            )

            seleccion = df_actividades.copy()

            if busqueda:
                mask = seleccion.apply(
                    lambda fila: fila.astype(str)
                    .str.contains(
                        busqueda,
                        case=False,
                        na=False,
                        regex=False,
                    )
                    .any(),
                    axis=1,
                )
                seleccion = seleccion[mask]

            if seleccion.empty:
                st.warning("No se encontraron registros.")
            else:
                ids_disponibles = seleccion["ID"].astype(str).tolist()

                def etiqueta_actividad(actividad_id):
                    fila_etiqueta = df_actividades[
                        df_actividades["ID"].astype(str) == str(actividad_id)
                    ].iloc[0]
                    return (
                        f"{actividad_id} — "
                        f"{valor_fila(fila_etiqueta, 'Título', 'Sin título')}"
                    )

                actividad_id = st.selectbox(
                    "Selecciona una actividad",
                    ids_disponibles,
                    format_func=etiqueta_actividad,
                )

                fila = df_actividades[
                    df_actividades["ID"].astype(str) == str(actividad_id)
                ].iloc[0].copy()

                prob_actuales = probatorios_de_actividad(
                    df_probatorios,
                    actividad_id,
                )

                st.markdown(f"### {actividad_id}")
                st.info(f"Esta actividad tiene **{len(prob_actuales)} probatorio(s)**.")

                with st.form(f"form_editar_{actividad_id}"):
                    col1, col2 = st.columns(2)

                    with col1:
                        anio = st.number_input(
                            "Año",
                            min_value=1900,
                            max_value=2100,
                            value=normalizar_anio(valor_fila(fila, "Año", datetime.now().year)),
                            step=1,
                        )

                        fecha = st.date_input(
                            "Fecha",
                            value=normalizar_fecha(valor_fila(fila, "Fecha", "")),
                        )

                        categoria_actual = str(
                            valor_fila(fila, "Categoría", CATEGORIAS[0])
                        )
                        indice_categoria = (
                            CATEGORIAS.index(categoria_actual)
                            if categoria_actual in CATEGORIAS
                            else 0
                        )

                        categoria = st.selectbox(
                            "Categoría",
                            CATEGORIAS,
                            index=indice_categoria,
                        )

                        rol = st.text_input(
                            "Rol",
                            value=str(valor_fila(fila, "Rol", "")),
                        )

                    with col2:
                        titulo = st.text_input(
                            "Título *",
                            value=str(valor_fila(fila, "Título", "")),
                        )

                        institucion = st.text_input(
                            "Institución",
                            value=str(valor_fila(fila, "Institución", "")),
                        )

                        lugar = st.text_input(
                            "Lugar / Sede",
                            value=str(valor_fila(fila, "Lugar", "")),
                        )

                        opciones_estado = [
                            "Verificado / En Drive",
                            "Pendiente de Escanear",
                            "En Trámite",
                        ]
                        estado_actual = str(
                            valor_fila(
                                fila,
                                "Estado_Probatorio",
                                "Verificado / En Drive",
                            )
                        )
                        indice_estado = (
                            opciones_estado.index(estado_actual)
                            if estado_actual in opciones_estado
                            else 0
                        )

                        estado = st.selectbox(
                            "Estado",
                            opciones_estado,
                            index=indice_estado,
                        )

                    incluir = st.radio(
                        "¿Incluir en CV?",
                        ["Sí", "No"],
                        horizontal=True,
                        index=(
                            0
                            if str(valor_fila(fila, "Incluir_en_CV", "No")).strip().lower()
                            == "sí"
                            else 1
                        ),
                    )

                    detalle_cv = st.text_area(
                        "Información complementaria para CV",
                        value=str(valor_fila(fila, "Detalle_CV", "")),
                        help=(
                            "Texto opcional que sí aparecerá en el CV generado. "
                            "Úsalo para datos curriculares relevantes que no tienen un campo propio, "
                            "como duración, nombre del evento general, volumen, número o contexto de la participación."
                        ),
                    )

                    notas = st.text_area(
                        "Notas / Observaciones",
                        value=str(valor_fila(fila, "Notas_Observaciones", "")),
                        help=(
                            "Uso interno de la base. Estas notas NO se incorporan automáticamente al CV."
                        ),
                    )

                    st.markdown("### 📎 Probatorios actuales")

                    ids_a_eliminar = []

                    if prob_actuales.empty:
                        st.warning("Esta actividad todavía no tiene probatorios.")
                    else:
                        for _, p in prob_actuales.iterrows():
                            pid = str(p["ID_Probatorio"])
                            nombre = str(p.get("Nombre_Archivo", "") or "Sin nombre")
                            enlace = str(p.get("Enlace_Drive_Probatorio", "") or "")

                            st.markdown(f"**{pid} — {nombre}**")
                            if enlace:
                                st.markdown(f"[🔗 Abrir en Drive]({enlace})")

                            marcar = st.checkbox(
                                f"Eliminar {pid}",
                                key=f"del_{actividad_id}_{pid}",
                            )
                            if marcar:
                                ids_a_eliminar.append(pid)

                    st.markdown("### 🔄 Reemplazar un probatorio")

                    opciones_reemplazo = ["Ninguno"] + (
                        prob_actuales["ID_Probatorio"].astype(str).tolist()
                        if not prob_actuales.empty
                        else []
                    )

                    id_reemplazo = st.selectbox(
                        "Probatorio a reemplazar",
                        opciones_reemplazo,
                    )

                    archivo_reemplazo = st.file_uploader(
                        "Nuevo archivo para el probatorio seleccionado",
                        type=["pdf", "png", "jpg", "jpeg"],
                        accept_multiple_files=False,
                        key=f"replace_file_{actividad_id}",
                    )

                    st.markdown("### ➕ Agregar nuevos probatorios")

                    nuevos_archivos = st.file_uploader(
                        "Puedes seleccionar uno o varios archivos",
                        type=["pdf", "png", "jpg", "jpeg"],
                        accept_multiple_files=True,
                        key=f"edit_files_{actividad_id}",
                    )

                    guardar = st.form_submit_button(
                        "💾 Guardar cambios",
                        type="primary",
                    )

                if guardar:
                    if not titulo.strip():
                        st.error("El título es obligatorio.")
                        st.stop()

                    if archivo_reemplazo and id_reemplazo == "Ninguno":
                        st.error("Selecciona qué probatorio quieres reemplazar.")
                        st.stop()

                    if id_reemplazo != "Ninguno" and not archivo_reemplazo:
                        st.error("Seleccionaste un probatorio para reemplazar, pero falta el archivo nuevo.")
                        st.stop()

                    if id_reemplazo in ids_a_eliminar:
                        st.error("Un mismo probatorio no puede eliminarse y reemplazarse al mismo tiempo.")
                        st.stop()

                    with st.spinner("Actualizando actividad, probatorios y Drive..."):
                        df_act_nuevo = df_actividades.copy()
                        df_prob_nuevo = df_probatorios.copy()

                        # ------------------------------------
                        # 1. ELIMINAR PROBATORIOS MARCADOS
                        # ------------------------------------
                        for pid in ids_a_eliminar:
                            coincidencia = df_prob_nuevo[
                                df_prob_nuevo["ID_Probatorio"].astype(str) == str(pid)
                            ]
                            if not coincidencia.empty:
                                drive_id = str(
                                    coincidencia.iloc[0].get("ID_Drive_Probatorio", "") or ""
                                ).strip()
                                if drive_id:
                                    eliminar_archivo_drive(service, drive_id)

                                df_prob_nuevo = df_prob_nuevo[
                                    df_prob_nuevo["ID_Probatorio"].astype(str) != str(pid)
                                ].copy()

                        # ------------------------------------
                        # 2. REEMPLAZAR UN PROBATORIO
                        # Conserva ID_Probatorio, cambia archivo de Drive.
                        # ------------------------------------
                        if id_reemplazo != "Ninguno" and archivo_reemplazo:
                            mask_reemplazo = (
                                df_prob_nuevo["ID_Probatorio"].astype(str)
                                == str(id_reemplazo)
                            )

                            if not mask_reemplazo.any():
                                st.error("No se encontró el probatorio seleccionado para reemplazar.")
                                st.stop()

                            indice_prob = df_prob_nuevo[mask_reemplazo].index[0]
                            fila_prob = df_prob_nuevo.loc[indice_prob]
                            orden_prob = normalizar_anio(fila_prob.get("Orden", 1), default=1)
                            drive_id_anterior = str(
                                fila_prob.get("ID_Drive_Probatorio", "") or ""
                            ).strip()

                            extension = os.path.splitext(archivo_reemplazo.name)[1].lower()
                            nombre_nuevo = nombre_probatorio(
                                titulo,
                                anio,
                                categoria,
                                orden_prob,
                                extension,
                            )

                            carpeta = obtener_carpeta_probatorio(
                                service,
                                estructura_drive,
                                anio,
                                categoria,
                            )

                            resultado_reemplazo = subir_a_google_drive(
                                service,
                                nombre_nuevo,
                                archivo_reemplazo.getvalue(),
                                carpeta,
                            )

                            if not resultado_reemplazo or not resultado_reemplazo.get("id"):
                                st.error("No se pudo subir el archivo de reemplazo.")
                                st.stop()

                            # Solo enviamos el viejo a papelera cuando el nuevo ya existe.
                            if drive_id_anterior:
                                eliminar_archivo_drive(service, drive_id_anterior)

                            df_prob_nuevo.at[indice_prob, "Nombre_Archivo"] = resultado_reemplazo["nombre"]
                            df_prob_nuevo.at[indice_prob, "Enlace_Drive_Probatorio"] = resultado_reemplazo["enlace"]
                            df_prob_nuevo.at[indice_prob, "ID_Drive_Probatorio"] = resultado_reemplazo["id"]
                            df_prob_nuevo.at[indice_prob, "Tipo_Archivo"] = resultado_reemplazo["mimetype"]
                            df_prob_nuevo.at[indice_prob, "Fecha_Registro"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

                        # ------------------------------------
                        # 3. AGREGAR NUEVOS PROBATORIOS
                        # ------------------------------------
                        sub_actual = probatorios_de_actividad(df_prob_nuevo, actividad_id)
                        orden_inicial = len(sub_actual) + 1

                        filas_nuevas, _ = crear_filas_probatorios_subidos(
                            service,
                            estructura_drive,
                            nuevos_archivos,
                            df_prob_nuevo,
                            actividad_id,
                            titulo.strip(),
                            anio,
                            categoria,
                            orden_inicial,
                        )

                        if filas_nuevas is None:
                            st.error("No se pudieron subir todos los nuevos probatorios.")
                            st.stop()

                        if filas_nuevas:
                            df_prob_nuevo = pd.concat(
                                [df_prob_nuevo, pd.DataFrame(filas_nuevas)],
                                ignore_index=True,
                            )

                        # ------------------------------------
                        # 4. REORDENAR, MOVER, RENOMBRAR Y
                        #    ASEGURAR URL DE TODOS LOS ARCHIVOS
                        # ------------------------------------
                        df_prob_nuevo = sincronizar_nombres_y_carpeta(
                            service,
                            estructura_drive,
                            df_prob_nuevo,
                            actividad_id,
                            titulo.strip(),
                            anio,
                            categoria,
                        )

                        # ------------------------------------
                        # 5. ACTUALIZAR ACTIVIDAD
                        # ------------------------------------
                        mask_act = df_act_nuevo["ID"].astype(str) == str(actividad_id)
                        indice_act = df_act_nuevo[mask_act].index[0]

                        tiene_probatorios = not probatorios_de_actividad(
                            df_prob_nuevo,
                            actividad_id,
                        ).empty

                        df_act_nuevo.at[indice_act, "Año"] = anio
                        df_act_nuevo.at[indice_act, "Fecha"] = str(fecha)
                        df_act_nuevo.at[indice_act, "Categoría"] = categoria
                        df_act_nuevo.at[indice_act, "Rol"] = rol
                        df_act_nuevo.at[indice_act, "Título"] = titulo.strip()
                        df_act_nuevo.at[indice_act, "Institución"] = institucion
                        df_act_nuevo.at[indice_act, "Lugar"] = lugar
                        df_act_nuevo.at[indice_act, "Estado_Probatorio"] = (
                            "Verificado / En Drive" if tiene_probatorios else estado
                        )
                        df_act_nuevo.at[indice_act, "Incluir_en_CV"] = incluir
                        df_act_nuevo.at[indice_act, "Detalle_CV"] = detalle_cv
                        df_act_nuevo.at[indice_act, "Notas_Observaciones"] = notas

                        ok = actualizar_excel_drive(
                            service,
                            excel_id,
                            df_act_nuevo,
                            preparar_probatorios(df_prob_nuevo, service=service),
                        )

                    if ok:
                        st.success("✅ Actividad actualizada correctamente.")
                        st.rerun()

                # --------------------------------------------
                # ELIMINAR ACTIVIDAD COMPLETA
                # --------------------------------------------
                st.markdown("---")
                st.markdown("### 🗑️ Eliminar actividad completa")
                st.caption(
                    "Esto enviará a la papelera de Drive todos los probatorios asociados "
                    "y eliminará las filas de Actividades y Probatorios."
                )

                confirmar_eliminacion = st.checkbox(
                    f"Confirmo que quiero eliminar {actividad_id}",
                    key=f"confirm_delete_activity_{actividad_id}",
                )

                if st.button(
                    "🗑️ Eliminar actividad",
                    disabled=not confirmar_eliminacion,
                    key=f"delete_activity_{actividad_id}",
                ):
                    with st.spinner("Eliminando actividad y probatorios..."):
                        asociados = probatorios_de_actividad(df_probatorios, actividad_id)

                        for _, p in asociados.iterrows():
                            drive_id = str(p.get("ID_Drive_Probatorio", "") or "").strip()
                            if drive_id:
                                eliminar_archivo_drive(service, drive_id)

                        df_act_nuevo = df_actividades[
                            df_actividades["ID"].astype(str) != str(actividad_id)
                        ].copy()

                        df_prob_nuevo = df_probatorios[
                            df_probatorios["ID_Actividad"].astype(str) != str(actividad_id)
                        ].copy()

                        ok = actualizar_excel_drive(
                            service,
                            excel_id,
                            df_act_nuevo,
                            df_prob_nuevo,
                        )

                    if ok:
                        st.success(f"Actividad {actividad_id} eliminada.")
                        st.rerun()

    # ========================================================
    # NUEVA ACTIVIDAD
    # ========================================================

    with tab_nuevo:
        st.subheader("➕ Registrar nueva actividad")

        with st.form("form_nueva_actividad"):
            col1, col2 = st.columns(2)

            with col1:
                nuevo_id = st.text_input(
                    "ID",
                    value=siguiente_id_registro(df_actividades),
                )

                anio = st.number_input(
                    "Año",
                    min_value=1900,
                    max_value=2100,
                    value=datetime.now().year,
                    step=1,
                )

                fecha = st.date_input("Fecha")
                categoria = st.selectbox("Categoría", CATEGORIAS)
                rol = st.text_input("Rol")

            with col2:
                titulo = st.text_input("Título *")
                institucion = st.text_input("Institución")
                lugar = st.text_input("Lugar / Sede")

                estado = st.selectbox(
                    "Estado",
                    [
                        "Verificado / En Drive",
                        "Pendiente de Escanear",
                        "En Trámite",
                    ],
                )

                incluir = st.radio(
                    "¿Incluir en CV?",
                    ["Sí", "No"],
                    horizontal=True,
                )

            st.markdown("---")
            st.subheader("📎 Probatorios")

            archivos = st.file_uploader(
                "Selecciona uno o varios probatorios",
                type=["pdf", "png", "jpg", "jpeg"],
                accept_multiple_files=True,
            )

            detalle_cv = st.text_area(
                "Información complementaria para CV",
                help=(
                    "Texto opcional que sí aparecerá en el CV generado. "
                    "Úsalo para información curricular complementaria, no para notas internas."
                ),
            )

            notas = st.text_area(
                "Notas / Observaciones",
                help=(
                    "Uso interno de la base. Estas notas NO se incorporan automáticamente al CV."
                ),
            )

            guardar_nueva = st.form_submit_button(
                "💾 Guardar actividad",
                type="primary",
            )

        if guardar_nueva:
            nuevo_id = nuevo_id.strip()

            if not nuevo_id:
                st.error("El ID de la actividad es obligatorio.")
                st.stop()

            if not titulo.strip():
                st.error("El título es obligatorio.")
                st.stop()

            if (
                df_actividades["ID"].astype(str).str.strip().str.lower()
                == nuevo_id.lower()
            ).any():
                st.error(f"Ya existe una actividad con el ID {nuevo_id}.")
                st.stop()

            with st.spinner("Subiendo actividad y probatorios..."):
                filas_probatorios, ids_drive_subidos = crear_filas_probatorios_subidos(
                    service,
                    estructura_drive,
                    archivos,
                    df_probatorios,
                    nuevo_id,
                    titulo.strip(),
                    anio,
                    categoria,
                    1,
                )

                if filas_probatorios is None:
                    st.error("No se pudieron subir todos los probatorios.")
                    st.stop()

                nueva_fila = {
                    "ID": nuevo_id,
                    "Año": anio,
                    "Fecha": str(fecha),
                    "Categoría": categoria,
                    "Rol": rol,
                    "Título": titulo.strip(),
                    "Institución": institucion,
                    "Lugar": lugar,
                    "Estado_Probatorio": (
                        "Verificado / En Drive" if filas_probatorios else estado
                    ),
                    "Incluir_en_CV": incluir,
                    "Detalle_CV": detalle_cv,
                    "Notas_Observaciones": notas,
                }

                df_act_nuevo = pd.concat(
                    [
                        df_actividades,
                        pd.DataFrame([nueva_fila], columns=COLUMNAS_ACTIVIDADES),
                    ],
                    ignore_index=True,
                )

                df_prob_nuevo = df_probatorios.copy()
                if filas_probatorios:
                    df_prob_nuevo = pd.concat(
                        [df_prob_nuevo, pd.DataFrame(filas_probatorios)],
                        ignore_index=True,
                    )

                ok = actualizar_excel_drive(
                    service,
                    excel_id,
                    preparar_actividades(df_act_nuevo),
                    preparar_probatorios(df_prob_nuevo, service=service),
                )

                if not ok:
                    # Si el Excel falló, evitamos dejar archivos recién subidos huérfanos.
                    for drive_id in ids_drive_subidos:
                        eliminar_archivo_drive(service, drive_id)
                    st.stop()

            st.success(f"✅ Actividad '{titulo}' registrada.")
            st.info(
                f"Se guardaron **{len(filas_probatorios)} probatorio(s)** en Drive y "
                "cada uno quedó registrado en la hoja Probatorios con su URL."
            )
            st.rerun()

    # ========================================================
    # PAQUETES ZIP
    # ========================================================

    with tab_paquetes:
        st.subheader("📦 Generador de paquetes de probatorios")
        st.write(
            "Selecciona actividades o utiliza filtros para construir un ZIP con "
            "todos sus probatorios."
        )

        st.markdown("### 🎯 Filtrar actividades")
        col1, col2, col3 = st.columns(3)

        with col1:
            anios_zip = sorted(
                pd.to_numeric(df_actividades["Año"], errors="coerce")
                .dropna()
                .astype(int)
                .unique()
                .tolist(),
                reverse=True,
            )
            filtro_zip_anio = st.selectbox(
                "Año",
                ["Todos"] + anios_zip,
                key="zip_anio",
            )

        with col2:
            filtro_zip_categoria = st.selectbox(
                "Categoría",
                ["Todas"] + CATEGORIAS,
                key="zip_categoria",
            )

        with col3:
            filtro_zip_texto = st.text_input("Buscar", key="zip_texto")

        df_zip = df_actividades.copy()

        if filtro_zip_anio != "Todos":
            df_zip = df_zip[
                pd.to_numeric(df_zip["Año"], errors="coerce")
                == int(filtro_zip_anio)
            ]

        if filtro_zip_categoria != "Todas":
            df_zip = df_zip[df_zip["Categoría"] == filtro_zip_categoria]

        if filtro_zip_texto:
            mask = df_zip.apply(
                lambda fila: fila.astype(str)
                .str.contains(
                    filtro_zip_texto,
                    case=False,
                    na=False,
                    regex=False,
                )
                .any(),
                axis=1,
            )
            df_zip = df_zip[mask]

        st.write(f"**{len(df_zip)} actividades encontradas.**")

        if not df_zip.empty:
            opciones = df_zip["ID"].astype(str).tolist()

            seleccionar_todas = st.checkbox(
                "☑️ Seleccionar todas las actividades filtradas",
                key="zip_todas",
            )

            if seleccionar_todas:
                seleccion_ids = opciones
                st.caption(f"Se incluirán las {len(opciones)} actividades filtradas.")
            else:
                seleccion_ids = st.multiselect(
                    "Selecciona las actividades",
                    opciones,
                    key="zip_multiselect",
                    format_func=lambda x: (
                        f"{x} — "
                        f"{valor_fila(df_actividades[df_actividades['ID'].astype(str) == str(x)].iloc[0], 'Título', 'Sin título')}"
                    ),
                )

            if seleccion_ids:
                df_seleccion_zip = df_actividades[
                    df_actividades["ID"].astype(str).isin(
                        [str(x) for x in seleccion_ids]
                    )
                ].copy()
            else:
                df_seleccion_zip = pd.DataFrame(columns=df_actividades.columns)

            st.write(f"**{len(df_seleccion_zip)} actividades seleccionadas.**")

            if st.button(
                "📦 Generar ZIP",
                type="primary",
                disabled=df_seleccion_zip.empty,
            ):
                with st.spinner("Descargando probatorios y construyendo ZIP..."):
                    zip_bytes, cantidad = crear_zip_probatorios(
                        service,
                        df_seleccion_zip,
                        df_probatorios,
                    )

                if cantidad == 0:
                    st.error("No se encontraron archivos válidos para estas actividades.")
                else:
                    fecha_zip = datetime.now().strftime("%Y%m%d_%H%M")
                    nombre_zip = f"Paquete_Probatorios_{fecha_zip}.zip"

                    st.success(f"✅ ZIP generado con **{cantidad} archivos**.")

                    st.download_button(
                        "📥 Descargar ZIP",
                        data=zip_bytes,
                        file_name=nombre_zip,
                        mime="application/zip",
                        type="primary",
                    )
        else:
            st.info("No hay actividades que coincidan con los filtros.")

    # ========================================================
    # GENERAR CV
    # ========================================================

    with tab_cv:
        st.subheader("📄 Generar CV en Word")

        df_cv = df_actividades[
            df_actividades["Incluir_en_CV"]
            .astype(str)
            .str.strip()
            .str.lower()
            == "sí"
        ].copy()

        st.info(
            f"Actualmente hay **{len(df_cv)} actividades** marcadas para el CV."
        )

        if not df_cv.empty:
            documento = crear_cv_word(df_actividades)

            st.download_button(
                "📥 Descargar CV (.docx)",
                data=documento,
                file_name="CV_Dra_Maria_Griselda_Gunther.docx",
                mime=(
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
                type="primary",
            )

            st.markdown("### 👁️ Actividades incluidas")
            st.dataframe(
                df_cv[
                    [
                        "ID",
                        "Año",
                        "Fecha",
                        "Categoría",
                        "Rol",
                        "Título",
                        "Institución",
                        "Detalle_CV",
                    ]
                ],
                use_container_width=True,
                hide_index=True,
            )
        else:
            st.warning("No hay actividades marcadas para incluir en el CV.")
